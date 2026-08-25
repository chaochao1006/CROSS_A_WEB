from __future__ import annotations

import argparse
import contextlib
import json
import logging
import math
import os
import re
import sys
import time
import traceback
from dataclasses import dataclass, field
from datetime import date, datetime, time as dtime, timedelta
from io import StringIO
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from zoneinfo import ZoneInfo

os.environ.setdefault("PANDAS_USE_NUMEXPR", "0")
os.environ.setdefault("PANDAS_USE_BOTTLENECK", "0")

with contextlib.redirect_stderr(StringIO()):
    import numpy as np
    import pandas as pd

try:
    import akshare as ak
except Exception:
    ak = None

try:
    from docx import Document
    from docx.shared import Pt
except Exception:
    Document = None
    Pt = None

try:
    from openpyxl import Workbook
    from openpyxl.chart import LineChart, Reference
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter
except Exception:
    Workbook = None
    LineChart = None
    Reference = None
    Alignment = None
    Font = None
    PatternFill = None
    get_column_letter = None


OUTPUT_DIR = Path(__file__).resolve().parent
EXCEL_TRACKING_DIR = OUTPUT_DIR / "EXCLE"
GOOGLE_SHEET_ID = ""
GOOGLE_WORKSHEET_NAME = "A-CROSS"
GOOGLE_SERVICE_ACCOUNT_ENV = "GOOGLE_SERVICE_ACCOUNT_JSON"
GOOGLE_SERVICE_ACCOUNT_FILE = OUTPUT_DIR / "google_service_account.json"
TRACKING_DAYS = 30
TICKERS = [
    "000725", "000977", "603083", "603629", "601869", "600105", "002050",
    "600580", "002472", "001309", "600584", "000021", "600703", "002156",
    "000988", "001270", "600879", "603596", "002405", "300408", "605376",
    "300623", "002371", "600667", "002384", "002428", "002409", "002463",
    "000636", "301217", "002636", "300604", "002879", "600183", "002353",
    "300383", "300666", "300351", "603019", "300657", "600378", "002820",
    "002965", "603666", "000880", "603918", "002583", "600126", "600843",
    "002131", "600186", "603927", "601116", "603108", "688246", "002044",
    "002085", "001696", "002065", "600602", "300115", "002696", "002350",
    "600362", "002757", "002837", "603086", "600089", "002498",
    "002792", "002465", "002361", "002268", "002119", "002150", "002407",
    "002929", "002709", "605336", "603881", "603667", "301366", "605123",
    "601991", "600396", "000692", "002886", "600143", "000969",
    "002741", "300308", "002185", "603267", "000539", "601686",
    "603890", "001314", "603936", "603005", "600732", "000970", "002971",
    "605589", "003026", "601208", "002637", "688825",
    "603716", "688836", "688826", "000657", "600498", "600011", "600578",
    "002506", "603601", "002364", "603618", "600522", "603678", "002951",
    "603011", "002436", "002579", "002552", "002173", "603045", "605358",
    "600487", "002916", "600641", "603773", "600206", "603986", "603186",
]

AK_ADJUST = "qfq"
AK_LOOKBACK_DAYS = 800
AK_FETCH_RETRIES = 3
TRIGGER_SCORE = 60
REQUEST_TIMEOUT = 15
SLEEP_BETWEEN_TICKERS = 0.2
MACD_FAST = 12
MACD_SLOW = 26
MACD_SIGNAL = 9
KDJ_PERIOD = 9
K_SMOOTH = 3
D_SMOOTH = 3


@dataclass
class StockResult:
    ticker: str
    company_name: str = ""
    df: Optional[pd.DataFrame] = None
    data_source: str = ""
    data_error: str = ""
    data_date: Optional[pd.Timestamp] = None
    latest_close: Optional[float] = None
    latest_change_pct: Optional[float] = None
    volume_ratio: Optional[float] = None
    score_parts: Dict[str, float] = field(default_factory=dict)
    total_score: float = 0.0
    signal_level: str = "不触发"
    macd_status: str = "无金叉"
    kdj_status: str = "无金叉"
    rsi_status: str = "无金叉"
    near_cross_notes: List[str] = field(default_factory=list)
    technical_summary: str = ""
    reason_categories: List[str] = field(default_factory=list)
    confidence: str = "低"
    reason_summary: str = ""
    risks: List[str] = field(default_factory=list)


def setup_logging(output_dir: Path) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    log_path = output_dir / "A-CROSS.log"
    logging.getLogger("peewee").setLevel(logging.CRITICAL)
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(levelname)s] %(message)s",
        handlers=[
            logging.FileHandler(log_path, encoding="utf-8"),
            logging.StreamHandler(sys.stdout),
        ],
    )


def china_now() -> datetime:
    return datetime.now(ZoneInfo("Asia/Shanghai"))


def fmt_date(value) -> str:
    if value is None or pd.isna(value):
        return "N/A"
    if isinstance(value, pd.Timestamp):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, date):
        return value.strftime("%Y-%m-%d")
    return str(value)


def safe_float(value) -> Optional[float]:
    try:
        if value is None or pd.isna(value):
            return None
        value = float(value)
        if not math.isfinite(value):
            return None
        return value
    except Exception:
        return None


def fmt_num(value, digits: int = 2) -> str:
    value = safe_float(value)
    if value is None:
        return "N/A"
    return f"{value:.{digits}f}"


def fmt_price(value) -> str:
    value = safe_float(value)
    if value is None:
        return "N/A"
    return f"¥{value:.2f}"


def clean_text(text: str) -> str:
    text = re.sub(r"<[^>]+>", " ", str(text or ""))
    return " ".join(text.split())


def display_stock(ticker: str, company_name: str = "") -> str:
    name = clean_text(company_name)
    return f"{ticker}（{name}）" if name and name != ticker else ticker


def split_trigger_record(line: str) -> Tuple[str, str, str]:
    line = line.strip()
    if "\t" in line:
        parts = line.split("\t")
        date_text = parts[0].strip() if parts else ""
        symbol_text = parts[1].strip() if len(parts) >= 2 else ""
        score_text = parts[2].strip() if len(parts) >= 3 else ""
        return date_text, symbol_text, score_text
    parts = line.split()
    date_text = parts[0].strip() if parts else ""
    symbol_text = parts[1].strip() if len(parts) >= 2 else ""
    score_text = parts[2].strip() if len(parts) >= 3 else ""
    return date_text, symbol_text, score_text


def parse_stock_symbol_text(symbol_text: str) -> Tuple[str, str]:
    text = clean_text(symbol_text)
    match = re.match(r"^(?P<ticker>\d{6})(?:[（(](?P<name>.*?)[）)])?", text)
    if not match:
        return text, ""
    ticker = match.group("ticker")
    name = clean_text(match.group("name") or "")
    if re.fullmatch(r"第\d+次", name):
        name = ""
    return ticker, name


def safe_filename_part(text: str) -> str:
    safe = clean_text(text)
    safe = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", safe)
    safe = re.sub(r"\s+", "_", safe).strip(" ._")
    return safe[:80]


_A_STOCK_NAME_MAP: Optional[Dict[str, str]] = None
_A_TRADE_DATES: Optional[List[date]] = None

STATIC_A_STOCK_NAMES = {
    "000725": "京东方A", "000977": "浪潮信息", "603083": "剑桥科技", "603629": "利通电子",
    "601869": "长飞光纤", "600105": "永鼎股份", "002050": "三花智控", "600580": "卧龙电驱",
    "002472": "双环传动", "001309": "德明利", "600584": "长电科技", "000021": "深科技",
    "600703": "三安光电", "002156": "通富微电", "000988": "华工科技", "001270": "铖昌科技",
    "600879": "航天电子", "603596": "伯特利", "002405": "四维图新", "300408": "三环集团",
    "605376": "博迁新材", "300623": "捷捷微电", "002371": "北方华创", "600667": "太极实业",
    "002384": "东山精密", "002428": "云南锗业", "002409": "雅克科技", "002463": "沪电股份",
    "000636": "风华高科", "301217": "铜冠铜箔", "002636": "金安国纪", "300604": "长川科技",
    "002879": "长缆科技", "600183": "生益科技", "002353": "杰瑞股份", "300383": "光环新网",
    "300666": "江丰电子", "300351": "永贵电器", "603019": "中科曙光", "300657": "弘信电子",
    "600378": "昊华科技", "002820": "桂发祥", "002965": "祥鑫科技", "603666": "亿嘉和",
    "000880": "潍柴重机", "603918": "金桥信息", "002583": "海能达", "600126": "杭钢股份",
    "600843": "上工申贝", "002131": "利欧股份", "600186": "莲花控股", "603927": "中科软",
    "601116": "三江购物", "603108": "润达医疗", "688246": "嘉和美康", "002044": "美年健康",
    "002085": "万丰奥威", "001696": "宗申动力", "002065": "东华软件", "600602": "云赛智联",
    "300115": "长盈精密", "002696": "百洋股份", "002350": "北京科锐", "601919": "中远海控",
    "600362": "江西铜业", "002757": "南兴股份", "002837": "英维克", "603086": "先达股份",
    "600089": "特变电工", "002498": "汉缆股份", "002792": "通宇通讯", "002465": "海格通信",
    "002361": "神剑股份", "002268": "电科网安", "002119": "康强电子", "002150": "正泰电源",
    "002407": "多氟多", "002929": "润建股份", "002709": "天赐材料", "605336": "*ST帅电",
    "603881": "数据港", "603667": "五洲新春", "301366": "一博科技", "605123": "派克新材",
    "601918": "新集能源", "601991": "大唐发电", "600396": "华电辽能", "000692": "惠天热电",
    "002886": "沃特股份", "600143": "金发科技", "000969": "安泰科技", "002741": "光华科技",
    "300308": "中际旭创", "002185": "华天科技", "603267": "鸿远电子", "000539": "粤电力A",
    "601686": "友发集团", "601101": "昊华能源", "603890": "春秋电子", "001314": "亿道信息",
    "603936": "博敏电子", "603005": "晶方科技", "600732": "爱旭股份", "000970": "中科三环",
    "002971": "和远气体", "605589": "圣泉集团", "003026": "中晶科技", "601208": "东材科技",
    "002637": "赞宇科技", "002165": "红宝丽", "688825": "长鑫科技", "601088": "中国神华",
    "603716": "塞力医疗", "688836": "C宇树", "688826": "C频准", "000657": "中钨高新",
    "600498": "烽火通信", "600011": "华能国际", "600578": "京能电力", "002506": "协鑫集成",
    "603601": "再升科技", "002364": "中恒电气", "603618": "杭电股份", "600522": "中天科技",
    "603678": "火炬电子", "002951": "金时科技", "603011": "合锻智能", "002436": "兴森科技",
    "002579": "中京电子", "002552": "宝鼎科技", "002173": "创新医疗", "603045": "福达合金",
    "605358": "立昂微", "600487": "亨通光电", "002916": "深南电路", "600641": "先导基电",
    "603773": "沃格光电", "600206": "有研新材", "603986": "兆易创新", "603186": "华正新材",
}


def get_a_stock_name_map() -> Dict[str, str]:
    global _A_STOCK_NAME_MAP
    if _A_STOCK_NAME_MAP is not None:
        return _A_STOCK_NAME_MAP
    _A_STOCK_NAME_MAP = {}
    if ak is None:
        return _A_STOCK_NAME_MAP
    for func_name in ("stock_info_a_code_name", "stock_zh_a_spot_em"):
        try:
            df = getattr(ak, func_name)()
            code_col = "code" if "code" in df.columns else "代码"
            name_col = "name" if "name" in df.columns else "名称"
            if code_col in df.columns and name_col in df.columns:
                codes = df[code_col].astype(str).str.zfill(6)
                names = df[name_col].astype(str).map(clean_text)
                _A_STOCK_NAME_MAP.update(dict(zip(codes, names)))
                if _A_STOCK_NAME_MAP:
                    break
        except Exception:
            continue
    return _A_STOCK_NAME_MAP


def get_company_name(ticker: str) -> str:
    code = ticker.zfill(6)
    return STATIC_A_STOCK_NAMES.get(code) or get_a_stock_name_map().get(code, ticker)


def get_a_trade_dates() -> List[date]:
    global _A_TRADE_DATES
    if _A_TRADE_DATES is not None:
        return _A_TRADE_DATES
    _A_TRADE_DATES = []
    if ak is None:
        return _A_TRADE_DATES
    try:
        calendar = ak.tool_trade_date_hist_sina()
        date_col = "trade_date" if "trade_date" in calendar.columns else calendar.columns[0]
        _A_TRADE_DATES = sorted(pd.to_datetime(calendar[date_col]).dt.date)
    except Exception:
        _A_TRADE_DATES = []
    return _A_TRADE_DATES


def previous_weekday(day: date) -> date:
    previous = day - timedelta(days=1)
    while previous.weekday() >= 5:
        previous -= timedelta(days=1)
    return previous


def latest_official_trading_day(now_cn: Optional[datetime] = None) -> date:
    now_cn = now_cn or china_now()
    today_cn = now_cn.date()
    trading_dates = get_a_trade_dates()
    previous_dates = [d for d in trading_dates if d < today_cn]
    if previous_dates:
        return previous_dates[-1]
    return previous_weekday(today_cn)


def ensure_ohlcv(df: pd.DataFrame) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame()
    df = df.copy()
    if isinstance(df.columns, pd.MultiIndex):
        df.columns = [c[0] if isinstance(c, tuple) else c for c in df.columns]
    rename_map = {}
    for col in df.columns:
        key = str(col).strip().lower()
        if key in {"open", "high", "low", "close", "volume"}:
            rename_map[col] = key.capitalize()
        elif key in {"adj close", "adjusted close"}:
            rename_map[col] = "Adjusted Close"
    df = df.rename(columns=rename_map)
    required = ["Open", "High", "Low", "Close", "Volume"]
    if any(c not in df.columns for c in required):
        return pd.DataFrame()
    if "Adjusted Close" not in df.columns:
        df["Adjusted Close"] = df["Close"]
    df = df[["Open", "High", "Low", "Close", "Adjusted Close", "Volume"]].copy()
    df.index = pd.to_datetime(df.index)
    if getattr(df.index, "tz", None) is not None:
        df.index = df.index.tz_convert("Asia/Shanghai").tz_localize(None)
    df.index = df.index.normalize()
    df = df.sort_index()
    df = df[~df.index.duplicated(keep="last")]
    for col in df.columns:
        df[col] = pd.to_numeric(df[col], errors="coerce")
    df = df.dropna(subset=["Open", "High", "Low", "Close", "Volume"])
    df = df[df["Close"] > 0]
    return df


def normalize_akshare_hist(df: pd.DataFrame) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame()
    df = df.copy()
    rename_map = {
        "日期": "Date",
        "date": "Date",
        "开盘": "Open",
        "open": "Open",
        "最高": "High",
        "high": "High",
        "最低": "Low",
        "low": "Low",
        "收盘": "Close",
        "close": "Close",
        "成交量": "Volume",
        "volume": "Volume",
        "amount": "Volume",
    }
    df = df.rename(columns={k: v for k, v in rename_map.items() if k in df.columns})
    required = ["Date", "Open", "High", "Low", "Close", "Volume"]
    if any(col not in df.columns for col in required):
        return pd.DataFrame()
    df["Adjusted Close"] = df["Close"]
    df = df.set_index("Date")
    return ensure_ohlcv(df)


def market_prefixed_symbol(ticker: str) -> str:
    code = ticker.zfill(6)
    prefix = "sh" if code.startswith(("5", "6", "9")) else "sz"
    return f"{prefix}{code}"


def fetch_akshare_em(ticker: str, start_date: str, end_date: str) -> Tuple[pd.DataFrame, str]:
    with contextlib.redirect_stdout(StringIO()), contextlib.redirect_stderr(StringIO()):
        df = ak.stock_zh_a_hist(
            symbol=ticker.zfill(6),
            period="daily",
            start_date=start_date,
            end_date=end_date,
            adjust=AK_ADJUST,
        )
    return normalize_akshare_hist(df), "AKShare stock_zh_a_hist"


def fetch_akshare_tx(ticker: str, start_date: str, end_date: str) -> Tuple[pd.DataFrame, str]:
    with contextlib.redirect_stdout(StringIO()), contextlib.redirect_stderr(StringIO()):
        df = ak.stock_zh_a_hist_tx(
            symbol=market_prefixed_symbol(ticker),
            start_date=start_date,
            end_date=end_date,
            adjust=AK_ADJUST,
            timeout=15,
        )
    return normalize_akshare_hist(df), "AKShare stock_zh_a_hist_tx"


def fetch_akshare_sina(ticker: str, start_date: str, end_date: str) -> Tuple[pd.DataFrame, str]:
    with contextlib.redirect_stdout(StringIO()), contextlib.redirect_stderr(StringIO()):
        df = ak.stock_zh_a_daily(
            symbol=market_prefixed_symbol(ticker),
            start_date=start_date,
            end_date=end_date,
            adjust=AK_ADJUST,
        )
    return normalize_akshare_hist(df), "AKShare stock_zh_a_daily"


def fetch_ohlcv(ticker: str, data_day: date) -> Tuple[pd.DataFrame, str, str]:
    errors = []
    if ak is None:
        return pd.DataFrame(), "", "未安装或无法导入 AKShare；请先运行 pip install akshare"
    end_date = data_day.strftime("%Y%m%d")
    start_date = (data_day - timedelta(days=AK_LOOKBACK_DAYS)).strftime("%Y%m%d")
    sources = (fetch_akshare_em, fetch_akshare_tx, fetch_akshare_sina)
    for source_func in sources:
        for attempt in range(1, AK_FETCH_RETRIES + 1):
            try:
                df, source = source_func(ticker, start_date, end_date)
                if df.empty:
                    errors.append(f"{source_func.__name__} 第 {attempt} 次：返回空数据或缺少 A 股日线 OHLCV 字段")
                    continue
                df = df[df.index.date <= data_day]
                if df.empty:
                    errors.append(f"{source}: 没有 {data_day} 及以前的正式收盘数据")
                    continue
                return df, f"{source}(adjust={AK_ADJUST or 'none'})", ""
            except Exception as exc:
                errors.append(f"{source_func.__name__} 第 {attempt} 次：{exc}")
            if attempt < AK_FETCH_RETRIES:
                time.sleep(1.5 * attempt)
    return pd.DataFrame(), "", "AKShare 异常：" + "；".join(errors[-3:])


def wilder_rsi(series: pd.Series, period: int) -> pd.Series:
    delta = series.diff()
    gain = delta.clip(lower=0)
    loss = -delta.clip(upper=0)
    avg_gain = gain.ewm(alpha=1 / period, adjust=False, min_periods=period).mean()
    avg_loss = loss.ewm(alpha=1 / period, adjust=False, min_periods=period).mean()
    rs = avg_gain / avg_loss.replace(0, np.nan)
    rsi = 100 - (100 / (1 + rs))
    return rsi.fillna(50)


def add_indicators(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    price = df["Adjusted Close"].where(df["Adjusted Close"].notna(), df["Close"])
    df["TrendClose"] = price
    df["MA20"] = price.rolling(20).mean()
    df["MA50"] = price.rolling(50).mean()
    df["Prev20High"] = df["High"].shift(1).rolling(20).max()
    df["VolumeMA20"] = df["Volume"].rolling(20).mean()
    df["VolumeRatio"] = df["Volume"] / df["VolumeMA20"].replace(0, np.nan)
    df["DailyChangePct"] = (price / price.shift(1) - 1) * 100
    daily_range = (df["High"] - df["Low"]).replace(0, np.nan)
    df["CloseLocation"] = ((df["Close"] - df["Low"]) / daily_range).fillna(0.5)

    ema_fast = price.ewm(span=MACD_FAST, adjust=False).mean()
    ema_slow = price.ewm(span=MACD_SLOW, adjust=False).mean()
    df["DIF"] = ema_fast - ema_slow
    df["DEA"] = df["DIF"].ewm(span=MACD_SIGNAL, adjust=False).mean()
    df["MACD_HIST"] = df["DIF"] - df["DEA"]

    low_n = df["Low"].rolling(KDJ_PERIOD).min()
    high_n = df["High"].rolling(KDJ_PERIOD).max()
    denom = (high_n - low_n).replace(0, np.nan)
    rsv = ((df["Close"] - low_n) / denom * 100).clip(0, 100).fillna(50)
    df["K"] = rsv.ewm(alpha=1 / K_SMOOTH, adjust=False).mean()
    df["D"] = df["K"].ewm(alpha=1 / D_SMOOTH, adjust=False).mean()
    df["J"] = 3 * df["K"] - 2 * df["D"]

    df["RSI6"] = wilder_rsi(price, 6)
    df["RSI12"] = wilder_rsi(price, 12)
    df["RSI24"] = wilder_rsi(price, 24)
    return df


def is_cross_up(fast: pd.Series, slow: pd.Series, offset: int = 0) -> bool:
    if len(fast) < offset + 2 or len(slow) < offset + 2:
        return False
    today = -1 - offset
    yesterday = -2 - offset
    return bool(fast.iloc[today] > slow.iloc[today] and fast.iloc[yesterday] <= slow.iloc[yesterday])


def crossed_within_two_days(df: pd.DataFrame, fast_col: str, slow_col: str) -> bool:
    fast = df[fast_col]
    slow = df[slow_col]
    today_cross = is_cross_up(fast, slow, 0)
    prev_cross = is_cross_up(fast, slow, 1)
    still_bullish = fast.iloc[-1] > slow.iloc[-1]
    return bool(still_bullish and (today_cross or prev_cross))


def is_gap_improving(df: pd.DataFrame, fast_col: str, slow_col: str, days: int = 3) -> bool:
    gap = (df[fast_col] - df[slow_col]).dropna().tail(days)
    if len(gap) < days:
        return False
    return bool(gap.iloc[-1] < 0 and (gap.diff().dropna() > 0).sum() >= days - 2)


def macd_zero_score(df: pd.DataFrame) -> float:
    row = df.iloc[-1]
    close = row["TrendClose"]
    if close <= 0:
        return 0
    ndif = row["DIF"] / close
    ndea = row["DEA"] / close
    if row["DIF"] > 0 and row["DEA"] > 0:
        return 10
    hist_norm = (df["DIF"] / df["TrendClose"].replace(0, np.nan)).dropna().abs()
    threshold = hist_norm.quantile(0.35) if len(hist_norm) >= 60 else 0.01
    if abs(ndif) <= threshold and abs(ndea) <= threshold:
        return 6
    return 0


def score_macd(df: pd.DataFrame, result: StockResult) -> float:
    row = df.iloc[-1]
    prev = df.iloc[-2]
    score = 0.0
    if crossed_within_two_days(df, "DIF", "DEA"):
        score += 10
        result.macd_status = "最近两个交易日内有效金叉"
    elif is_gap_improving(df, "DIF", "DEA"):
        result.macd_status = "即将金叉"
        result.near_cross_notes.append("MACD 即将金叉")
    score += macd_zero_score(df)

    hist = df["MACD_HIST"].tail(4)
    hist_diff = hist.diff().dropna()
    if len(hist) >= 4 and hist.iloc[-1] > 0 and (hist_diff.tail(3) > 0).all():
        score += 10
    elif len(hist) >= 3 and hist.iloc[-1] > 0 and (hist_diff.tail(2) > 0).all():
        score += 7
    elif row["MACD_HIST"] > prev["MACD_HIST"] or (row["MACD_HIST"] > 0 and prev["MACD_HIST"] <= 0):
        score += 4

    if row["DIF"] > prev["DIF"] and row["DEA"] > prev["DEA"]:
        score += 5
    return min(score, 35)


def score_kdj(df: pd.DataFrame, result: StockResult) -> float:
    row = df.iloc[-1]
    score = 0.0
    if crossed_within_two_days(df, "K", "D"):
        score += 10
        result.kdj_status = "最近两个交易日内有效金叉"
    elif is_gap_improving(df, "K", "D"):
        result.kdj_status = "即将金叉"
        result.near_cross_notes.append("KDJ 即将金叉")

    k, d, j = row["K"], row["D"], row["J"]
    center = (k + d) / 2
    if 30 <= center <= 65 or 40 <= center <= 70:
        score += 5
    elif 20 <= center <= 75:
        score += 3

    j_flags = [j > k, j > 50, j > df["J"].iloc[-2]]
    count = sum(bool(x) for x in j_flags)
    if count == 3:
        score += 5
    elif count == 2:
        score += 3
    elif count == 1:
        score += 1
    return min(score, 20)


def score_rsi(df: pd.DataFrame, result: StockResult) -> float:
    row = df.iloc[-1]
    prev = df.iloc[-2]
    score = 0.0
    rsi_crosses = []
    if is_cross_up(df["RSI6"], df["RSI12"], 0):
        rsi_crosses.append("RSI6上穿RSI12")
    if is_cross_up(df["RSI12"], df["RSI24"], 0):
        rsi_crosses.append("RSI12上穿RSI24")
    if rsi_crosses:
        result.rsi_status = "；".join(rsi_crosses)
    elif row["RSI6"] > row["RSI12"] > row["RSI24"]:
        result.rsi_status = "RSI多头排列"
    elif is_gap_improving(df, "RSI6", "RSI12") or is_gap_improving(df, "RSI12", "RSI24"):
        result.rsi_status = "即将金叉"
        result.near_cross_notes.append("RSI 即将金叉")

    if row["RSI6"] > row["RSI12"] > row["RSI24"]:
        score += 10
    if row["RSI12"] >= 60:
        score += 5
    elif row["RSI12"] >= 55:
        score += 4
    elif row["RSI12"] >= 50:
        score += 3
    if row["RSI24"] >= 50 and row["RSI24"] > prev["RSI24"]:
        score += 5
    elif row["RSI24"] >= 45 and row["RSI24"] > prev["RSI24"]:
        score += 3
    return min(score, 20)


def score_price_trend(df: pd.DataFrame) -> float:
    row = df.iloc[-1]
    score = 0.0
    above_ma20 = row["TrendClose"] > row["MA20"] if pd.notna(row["MA20"]) else False
    above_ma50 = row["TrendClose"] > row["MA50"] if pd.notna(row["MA50"]) else False
    if above_ma20 and above_ma50:
        score += 10
    elif above_ma20 or above_ma50:
        score += 5
    if pd.notna(row["Prev20High"]) and row["TrendClose"] > row["Prev20High"]:
        score += 5
    return min(score, 15)


def score_volume(df: pd.DataFrame) -> float:
    vr = safe_float(df.iloc[-1]["VolumeRatio"])
    if vr is None:
        return 0
    if vr >= 1.5:
        return 10
    if vr >= 1.3:
        return 8
    if vr >= 1.2:
        return 5
    return 0


def signal_level(score: float) -> str:
    if score >= 80:
        return "强势金叉"
    if score >= 71:
        return "较强金叉"
    if score >= 60:
        return "普通有效信号"
    return "不触发"


def build_technical_summary(result: StockResult) -> str:
    df = result.df
    if df is None or df.empty:
        return ""
    row = df.iloc[-1]
    points = []
    if result.macd_status != "无金叉":
        points.append(f"MACD{result.macd_status}，DIF/DEA为{fmt_num(row['DIF'], 4)}/{fmt_num(row['DEA'], 4)}")
    if result.kdj_status != "无金叉":
        points.append(f"KDJ{result.kdj_status}，K/D/J为{fmt_num(row['K'], 1)}/{fmt_num(row['D'], 1)}/{fmt_num(row['J'], 1)}")
    if result.rsi_status != "无金叉":
        points.append(f"RSI状态为{result.rsi_status}，RSI12为{fmt_num(row['RSI12'], 1)}")
    if row["TrendClose"] > row["MA20"] and row["TrendClose"] > row["MA50"]:
        points.append("股价站上MA20和MA50")
    elif row["TrendClose"] > row["MA20"] or row["TrendClose"] > row["MA50"]:
        points.append("股价站上部分均线")
    if pd.notna(row["Prev20High"]) and row["TrendClose"] > row["Prev20High"]:
        points.append("收盘价突破过去20日高点")
    elif pd.notna(row["Prev20High"]) and (row["Prev20High"] - row["TrendClose"]) / row["Prev20High"] <= 0.02:
        points.append("接近突破过去20日高点")
    if not points:
        points.append("技术指标尚未形成明显强势共振")
    return "；".join(points) + "。"


def analyze_reason(result: StockResult) -> None:
    result.reason_categories = ["技术性突破"] if result.total_score >= TRIGGER_SCORE else ["未触发"]
    result.confidence = "技术指标判断"
    result.reason_summary = "仅按技术指标触发。"


def build_risks(result: StockResult) -> List[str]:
    df = result.df
    if df is None or df.empty:
        return ["数据不足，风险无法完整评估。"]
    row = df.iloc[-1]
    risks = []
    if row["RSI6"] > 80:
        risks.append("RSI6高于80，短线存在过热风险。")
    if row["RSI12"] > 75:
        risks.append("RSI12高于75，需防短线震荡。")
    if row["K"] > 80 and row["D"] > 80:
        risks.append("KDJ处于高位，可能出现高位钝化或回落。")
    if pd.notna(row["MA20"]) and row["MA20"] > 0 and (row["TrendClose"] / row["MA20"] - 1) > 0.10:
        risks.append("股价远离MA20超过10%，短线追高风险较高。")
    if row["CloseLocation"] <= 0.3:
        risks.append("收盘位置靠近当日低位，存在冲高回落风险。")
    if result.volume_ratio is not None and result.volume_ratio < 1.2:
        risks.append("成交量未明显放大，需防信号确认不足。")
    if len([s for s in [result.macd_status, result.kdj_status, result.rsi_status] if s != "无金叉"]) <= 1:
        risks.append("仅少数指标确认，可能是假突破。")
    if not risks:
        risks.append("重点观察次日量能和指标延续性。")
    return risks[:2]


def analyze_ticker(ticker: str, data_day: date) -> StockResult:
    result = StockResult(ticker=ticker, company_name=get_company_name(ticker))
    df, source, err = fetch_ohlcv(ticker, data_day)
    if df.empty:
        result.data_error = err or "无法获取行情数据"
        logging.warning("%s 数据获取失败：%s", ticker, result.data_error)
        return result
    if len(df) < 80:
        result.data_error = "历史行情不足，无法稳定计算指标"
        logging.warning("%s 历史行情不足：%s 行", ticker, len(df))
        return result
    result.data_source = source
    df = add_indicators(df)
    df = df.dropna(subset=["MA20", "MA50", "DIF", "DEA", "MACD_HIST", "K", "D", "J", "RSI6", "RSI12", "RSI24"])
    if len(df) < 30:
        result.data_error = "指标有效数据不足"
        logging.warning("%s 指标有效数据不足", ticker)
        return result
    result.df = df
    row = df.iloc[-1]
    result.data_date = df.index[-1]
    if result.data_date.date() != data_day:
        result.data_error = f"最新行情日期为 {result.data_date.date()}，缺少 {data_day} 的正式收盘数据"
        logging.warning("%s 数据日期滞后：%s", ticker, result.data_error)
        return result
    result.latest_close = safe_float(row["TrendClose"])
    result.latest_change_pct = safe_float(row["DailyChangePct"])
    result.volume_ratio = safe_float(row["VolumeRatio"])

    result.score_parts["MACD"] = score_macd(df, result)
    result.score_parts["KDJ"] = score_kdj(df, result)
    result.score_parts["RSI"] = score_rsi(df, result)
    result.score_parts["PriceTrend"] = score_price_trend(df)
    result.score_parts["Volume"] = score_volume(df)
    result.total_score = max(0, min(100, sum(result.score_parts.values())))
    result.signal_level = signal_level(result.total_score)
    result.technical_summary = build_technical_summary(result)
    result.risks = build_risks(result)
    return result


def report_exists_for_day(output_dir: Path, data_day: date) -> Optional[Path]:
    path = output_dir / f"A-CROSS强势金叉监控报告_{data_day.strftime('%Y-%m-%d')}.docx"
    return path if path.exists() else None


def append_triggered_symbols_txt(output_dir: Path, report_date: date, triggered: List[StockResult]) -> List[Tuple[str, str, str]]:
    if not triggered:
        return []
    txt_path = output_dir / "CROSS.txt"
    existing = set()
    symbol_counts: Dict[str, int] = {}

    def parse_symbol_from_line(line: str) -> str:
        _, symbol_text, _ = split_trigger_record(line)
        if not symbol_text:
            return ""
        ticker, _ = parse_stock_symbol_text(symbol_text)
        return ticker

    def parse_date_symbol_key(line: str) -> str:
        date_text, _, _ = split_trigger_record(line)
        if not date_text:
            return line.strip()
        symbol_text = parse_symbol_from_line(line)
        return f"{date_text}\t{symbol_text}"

    if txt_path.exists():
        try:
            for raw_line in txt_path.read_text(encoding="utf-8-sig").splitlines():
                line = raw_line.strip()
                if not line:
                    continue
                existing.add(parse_date_symbol_key(line))
                symbol = parse_symbol_from_line(line)
                if symbol:
                    symbol_counts[symbol] = symbol_counts.get(symbol, 0) + 1
        except Exception as exc:
            logging.warning("读取 CROSS.txt 失败，将仅写入本次触发结果：%s", exc)
            existing = set()
            symbol_counts = {}

    new_records: List[Tuple[str, str, str]] = []
    date_text = report_date.strftime("%Y-%m-%d")
    for result in triggered:
        key = f"{date_text}\t{result.ticker}"
        if key in existing:
            continue
        next_count = symbol_counts.get(result.ticker, 0) + 1
        symbol_text = display_stock(result.ticker, result.company_name)
        if next_count > 1:
            symbol_text = f"{symbol_text}（第{next_count}次）"
        score_text = f"{result.total_score:.1f}分"
        new_records.append((date_text, symbol_text, score_text))
        existing.add(key)
        symbol_counts[result.ticker] = next_count

    if new_records:
        with txt_path.open("a", encoding="utf-8-sig") as f:
            for date_text, symbol_text, score_text in new_records:
                f.write(f"{date_text}\t{symbol_text}\t{score_text}\n")
        logging.info("已写入 CROSS.txt：%s", ", ".join(symbol for _, symbol, _ in new_records))
    else:
        logging.info("CROSS.txt 中已存在本交易日触发记录，未重复写入。")
    return new_records


def append_triggered_symbols_google_sheet(records: List[Tuple[str, str, str]]) -> None:
    if not records:
        return
    if not GOOGLE_SHEET_ID:
        logging.info("A-CROSS 未配置 Google Sheet ID，已跳过云端表格写入。")
        return
    credential_path = os.environ.get(GOOGLE_SERVICE_ACCOUNT_ENV, "").strip()
    if not credential_path and GOOGLE_SERVICE_ACCOUNT_FILE.exists():
        credential_path = str(GOOGLE_SERVICE_ACCOUNT_FILE)
    if not credential_path:
        logging.info(
            "未配置 Google Sheets 授权，已跳过云端表格写入。可设置环境变量 %s，或把服务账号 JSON 放到：%s",
            GOOGLE_SERVICE_ACCOUNT_ENV,
            GOOGLE_SERVICE_ACCOUNT_FILE,
        )
        return

    try:
        import gspread
    except Exception as exc:
        logging.warning("未安装或无法导入 gspread，已跳过 Google Sheets 写入：%s", exc)
        return

    try:
        client = gspread.service_account(filename=credential_path)
        spreadsheet = client.open_by_key(GOOGLE_SHEET_ID)
        try:
            worksheet = spreadsheet.worksheet(GOOGLE_WORKSHEET_NAME)
        except gspread.WorksheetNotFound:
            worksheet = spreadsheet.add_worksheet(title=GOOGLE_WORKSHEET_NAME, rows=1000, cols=3)
        if not worksheet.get_all_values():
            worksheet.append_row(["触发日期", "股票代码", "评分"], value_input_option="USER_ENTERED")
        worksheet.append_rows(
            [[date_text, symbol_text, score_text] for date_text, symbol_text, score_text in records],
            value_input_option="USER_ENTERED",
        )
        logging.info("已写入 Google Sheets：%s", ", ".join(symbol for _, symbol, _ in records))
    except Exception as exc:
        logging.warning("Google Sheets 写入失败，已保留本地 CROSS.txt 记录：%s", exc)


def parse_plain_ticker(symbol_text: str) -> str:
    symbol, _ = parse_stock_symbol_text(symbol_text)
    return symbol.strip().upper()


def read_cross_trigger_history(output_dir: Path) -> Dict[str, Tuple[str, List[date]]]:
    txt_path = output_dir / "CROSS.txt"
    history: Dict[str, List[date]] = {}
    names: Dict[str, str] = {}
    if not txt_path.exists():
        return history
    try:
        lines = txt_path.read_text(encoding="utf-8-sig").splitlines()
    except Exception as exc:
        logging.warning("读取 CROSS.txt 失败，无法更新 Excel 跟踪文件：%s", exc)
        return history

    for raw_line in lines:
        date_text, symbol_text, _ = split_trigger_record(raw_line)
        if not date_text or not symbol_text:
            continue
        try:
            trigger_date = date.fromisoformat(date_text)
        except ValueError:
            continue
        ticker, company_name = parse_stock_symbol_text(symbol_text)
        if not ticker:
            continue
        if company_name:
            names[ticker] = company_name
        history.setdefault(ticker, []).append(trigger_date)
    return {ticker: (names.get(ticker, ""), sorted(set(dates))) for ticker, dates in history.items()}


def safe_tracking_filename(ticker: str, company_name: str, episode_number: int) -> str:
    name = clean_text(company_name)
    stock_text = f"{ticker}_{name}" if name and name != ticker else ticker
    safe = safe_filename_part(stock_text)
    return f"{safe or 'UNKNOWN'}_{episode_number}"


def tracking_workbook_path(tracking_dir: Path, ticker: str, company_name: str, episode_number: int, trigger_date: date) -> Path:
    base_name = safe_tracking_filename(ticker, company_name, episode_number)
    dated_path = tracking_dir / f"{base_name}_{trigger_date.strftime('%Y%m%d')}.xlsx"
    return dated_path


def tracking_prices_from_df(df: pd.DataFrame, trigger_date: date, data_day: date) -> pd.DataFrame:
    df = df.copy()
    df["TrackClose"] = pd.to_numeric(df["Close"], errors="coerce")
    df["TrackDailyChangePct"] = df["TrackClose"].pct_change() * 100
    df = df[(df.index.date >= trigger_date) & (df.index.date <= data_day)].copy()
    if df.empty:
        return pd.DataFrame()
    df = df.head(TRACKING_DAYS)
    df = df.dropna(subset=["TrackClose"])
    return df


def is_trigger_inside_episode(df: pd.DataFrame, episode_start: date, trigger_date: date) -> bool:
    episode_window = df[(df.index.date >= episode_start) & (df.index.date <= trigger_date)]
    if episode_window.empty:
        return True
    return len(episode_window) <= TRACKING_DAYS


def build_tracking_episodes(df: pd.DataFrame, trigger_dates: List[date], data_day: date) -> List[date]:
    episodes: List[date] = []
    for trigger_date in sorted(d for d in set(trigger_dates) if d <= data_day):
        if not episodes:
            episodes.append(trigger_date)
            continue
        if not is_trigger_inside_episode(df, episodes[-1], trigger_date):
            episodes.append(trigger_date)
    return episodes


def write_tracking_workbook(path: Path, ticker: str, company_name: str, episode_number: int, trigger_date: date, data_day: date, df: pd.DataFrame) -> None:
    if Workbook is None:
        raise RuntimeError("未安装 openpyxl，无法生成 Excel 跟踪文件")
    path.parent.mkdir(parents=True, exist_ok=True)

    wb = Workbook()
    ws = wb.active
    ws.title = "30日跟踪"
    ws.freeze_panes = "A9"

    stock_label = display_stock(ticker, company_name)
    title = f"{stock_label}_第{episode_number}次触发后30个交易日股价跟踪"
    ws.merge_cells("A1:D1")
    ws["A1"] = title
    ws["A1"].font = Font(bold=True, size=14)
    ws["A1"].alignment = Alignment(horizontal="center")

    base_close = float(df.iloc[0]["TrackClose"])
    meta_rows = [
        ("股票代码", ticker),
        ("中文名", clean_text(company_name) or ticker),
        ("跟踪周期", f"第{episode_number}次"),
        ("首次触发日", trigger_date),
        ("触发日收盘价", base_close),
        ("更新到交易日", data_day),
    ]
    for idx, (label, value) in enumerate(meta_rows, start=2):
        ws.cell(row=idx, column=1, value=label)
        ws.cell(row=idx, column=2, value=value)
        ws.cell(row=idx, column=1).font = Font(bold=True)
    ws["B5"].number_format = "yyyy-mm-dd"
    ws["B6"].number_format = "0.00"
    ws["B7"].number_format = "yyyy-mm-dd"

    headers = ["交易日", "收盘价", "当日涨跌幅%", "相对触发日涨跌幅", "进度"]
    header_row = 8
    for col, header in enumerate(headers, start=1):
        cell = ws.cell(row=header_row, column=col, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F4E78")
        cell.alignment = Alignment(horizontal="center")

    first_data_row = header_row + 1
    for offset, (idx, row) in enumerate(df.iterrows(), start=0):
        excel_row = first_data_row + offset
        trade_date = idx.date() if hasattr(idx, "date") else idx
        ws.cell(row=excel_row, column=1, value=trade_date)
        ws.cell(row=excel_row, column=2, value=float(row["TrackClose"]))
        daily_change = row.get("TrackDailyChangePct")
        ws.cell(row=excel_row, column=3, value=None if pd.isna(daily_change) else float(daily_change))
        ws.cell(row=excel_row, column=4, value=f"=B{excel_row}/$B${first_data_row}-1")
        ws.cell(row=excel_row, column=5, value=f"第{offset + 1}/{TRACKING_DAYS}个交易日")
        ws.cell(row=excel_row, column=1).number_format = "yyyy-mm-dd"
        ws.cell(row=excel_row, column=2).number_format = "0.00"
        ws.cell(row=excel_row, column=3).number_format = "0.00"
        ws.cell(row=excel_row, column=4).number_format = "0.00%"

    last_data_row = first_data_row + len(df) - 1
    chart = LineChart()
    chart.title = f"{stock_label}_第{episode_number}次 相对触发日涨跌幅"
    chart.y_axis.title = "涨跌幅"
    chart.x_axis.title = "交易日"
    chart.y_axis.numFmt = "0%"
    chart.height = 10
    chart.width = 18
    data = Reference(ws, min_col=4, min_row=header_row, max_row=last_data_row)
    cats = Reference(ws, min_col=1, min_row=first_data_row, max_row=last_data_row)
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(cats)
    chart.legend = None
    ws.add_chart(chart, "F2")

    widths = {"A": 14, "B": 18, "C": 18, "D": 18}
    for col, width in widths.items():
        ws.column_dimensions[col].width = width
    for col in range(5, 15):
        ws.column_dimensions[get_column_letter(col)].width = 12

    wb.save(path)


def update_cross_tracking_excels(output_dir: Path, data_day: date) -> List[Path]:
    history = read_cross_trigger_history(output_dir)
    if not history:
        logging.info("CROSS.txt 暂无触发记录，未生成 Excel 跟踪文件。")
        return []
    if Workbook is None:
        logging.warning("未安装 openpyxl，跳过 Excel 跟踪文件生成。")
        return []

    updated_paths: List[Path] = []
    tracking_dir = EXCEL_TRACKING_DIR
    tracking_dir.mkdir(parents=True, exist_ok=True)

    for ticker, (company_name, trigger_dates) in sorted(history.items()):
        try:
            company_name = company_name or get_company_name(ticker)
            full_df, source, err = fetch_ohlcv(ticker, data_day)
            if full_df.empty:
                logging.warning("%s Excel 跟踪更新失败：%s", ticker, err or "无法获取行情数据")
                continue
            episodes = build_tracking_episodes(full_df, trigger_dates, data_day)
            for episode_number, trigger_date in enumerate(episodes, start=1):
                df = tracking_prices_from_df(full_df, trigger_date, data_day)
                if df.empty:
                    logging.warning("%s_%s Excel 跟踪更新失败：没有 %s 至 %s 的收盘数据", ticker, episode_number, trigger_date, data_day)
                    continue
                workbook_path = tracking_workbook_path(tracking_dir, ticker, company_name, episode_number, trigger_date)
                write_tracking_workbook(workbook_path, ticker, company_name, episode_number, trigger_date, data_day, df)
                updated_paths.append(workbook_path)
                logging.info(
                    "已更新 %s_%s 跟踪 Excel：%s，记录 %s/%s 个交易日。",
                    display_stock(ticker, company_name),
                    episode_number,
                    workbook_path,
                    len(df),
                    TRACKING_DAYS,
                )
        except Exception as exc:
            logging.warning("%s Excel 跟踪更新失败：%s", ticker, exc)
    return updated_paths


def write_word_report(path: Path, data_day: date, now_cn: datetime, results: List[StockResult], scanned: int, success_count: int) -> None:
    if Document is None:
        raise RuntimeError("未安装 python-docx，请运行：pip install python-docx")
    doc = Document()
    doc.styles["Normal"].font.name = "Microsoft YaHei"
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.add_heading(f"A-CROSS强势金叉监控报告_{data_day.strftime('%Y-%m-%d')}", level=1)

    doc.add_heading("一、报告概况", level=2)
    for text in [
        f"数据对应的 A 股交易日：{data_day.strftime('%Y-%m-%d')}",
        f"报告生成时间（北京时间）：{now_cn.strftime('%Y-%m-%d %H:%M:%S %Z')}",
        "取数规则：固定使用北京时间当前日期之前的最近一个 A 股交易日，不使用当天盘中或当天收盘数据。",
        f"数据源：AKShare（东方财富优先，腾讯/新浪备用），复权方式：{AK_ADJUST or '不复权'}",
        f"扫描股票总数：{scanned}",
        f"数据获取成功数量：{success_count}",
        f"评分达到{TRIGGER_SCORE}分及以上的触发股票数量：{len(results)}",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    if not results:
        doc.add_paragraph(f"本交易日没有股票达到{TRIGGER_SCORE}分触发阈值。")
        doc.add_heading("免责声明", level=2)
        doc.add_paragraph("本报告仅用于技术指标监控，不构成任何投资建议。技术指标可能存在滞后或误差，请自行判断投资风险。")
        doc.save(path)
        return

    doc.add_heading("二、触发股票汇总表", level=2)
    table = doc.add_table(rows=1, cols=8)
    table.style = "Table Grid"
    headers = ["股票代码", "中文名", "最新价", "当日涨跌幅", "总评分", "信号等级", "成交量倍数", "触发依据"]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
    for r in results:
        cells = table.add_row().cells
        values = [
            r.ticker,
            r.company_name,
            fmt_price(r.latest_close),
            f"{fmt_num(r.latest_change_pct, 2)}%",
            fmt_num(r.total_score, 1),
            r.signal_level,
            fmt_num(r.volume_ratio, 2),
            "、".join(r.reason_categories) if r.reason_categories else "技术性突破",
        ]
        for cell, value in zip(cells, values):
            cell.text = value

    doc.add_heading("三、触发股票简要分析", level=2)
    for r in results:
        doc.add_heading(f"{display_stock(r.ticker, r.company_name)}｜综合评分：{fmt_num(r.total_score, 1)}分｜{r.signal_level}", level=3)
        doc.add_paragraph(
            f"最新收盘价：{fmt_price(r.latest_close)}；当日涨跌幅：{fmt_num(r.latest_change_pct, 2)}%；成交量倍数：{fmt_num(r.volume_ratio, 2)}倍。"
            f"技术面：{r.technical_summary}"
            f"风险提示：{'；'.join(r.risks)}"
        )

    doc.add_heading("免责声明", level=2)
    doc.add_paragraph("本报告仅用于技术指标监控，不构成任何投资建议。技术指标可能存在滞后或误差，请自行判断投资风险。")
    doc.save(path)


def run(force: bool = False) -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    setup_logging(OUTPUT_DIR)
    now_cn = china_now()
    data_day = latest_official_trading_day(now_cn)
    out_path = OUTPUT_DIR / f"A-CROSS强势金叉监控报告_{data_day.strftime('%Y-%m-%d')}.docx"
    if out_path.exists() and not force:
        logging.info("报告已存在，未覆盖：%s。若需覆盖请使用 --force", out_path)
        update_cross_tracking_excels(OUTPUT_DIR, data_day)
        return out_path

    technical_results: List[StockResult] = []
    success_count = 0
    logging.info("开始扫描 %s 只股票，正式交易日：%s", len(TICKERS), data_day)
    for i, ticker in enumerate(TICKERS, 1):
        logging.info("[%s/%s] %s", i, len(TICKERS), ticker)
        try:
            result = analyze_ticker(ticker, data_day)
            if not result.data_error:
                success_count += 1
            technical_results.append(result)
            if result.total_score < TRIGGER_SCORE:
                logging.info("%s 综合评分为 %.1f，低于%s分，不触发。", ticker, result.total_score, TRIGGER_SCORE)
        except Exception as exc:
            logging.exception("%s 分析失败：%s", ticker, exc)
            technical_results.append(StockResult(ticker=ticker, company_name=get_company_name(ticker), data_error=str(exc)))
        time.sleep(SLEEP_BETWEEN_TICKERS)

    triggered = [r for r in technical_results if not r.data_error and r.total_score >= TRIGGER_SCORE]
    triggered.sort(key=lambda x: x.total_score, reverse=True)
    for r in triggered:
        logging.info("%s 达到 %.1f 分，仅记录技术信号。", r.ticker, r.total_score)
        analyze_reason(r)

    write_word_report(out_path, data_day, now_cn, triggered, len(TICKERS), success_count)
    new_records = append_triggered_symbols_txt(OUTPUT_DIR, data_day, triggered)
    append_triggered_symbols_google_sheet(new_records)
    update_cross_tracking_excels(OUTPUT_DIR, data_day)
    logging.info("报告生成完成：%s", out_path)
    return out_path


def main() -> None:
    parser = argparse.ArgumentParser(description="A股强势 MACD/KDJ/RSI 金叉监控程序")
    parser.add_argument("--force", action="store_true", help="若同一交易日报告已存在，则强制覆盖")
    args = parser.parse_args()
    try:
        path = run(force=args.force)
        print("")
        print("报告生成完成：")
        print(f"Word: {path}")
    except KeyboardInterrupt:
        print("用户中断。")
    except Exception as exc:
        print("运行失败：", exc)
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
